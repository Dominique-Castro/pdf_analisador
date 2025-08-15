import streamlit as st
import os
import numpy as np
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document
import io
from datetime import datetime
import re
import hashlib
import subprocess

# ========== CONFIGURAÇÃO INICIAL ========== #
st.set_page_config(
    page_title="Sistema de Análise Documental - BM/RS",
    page_icon="🛡️",
    layout="wide"
)

# Verificação do Poppler
def verificar_poppler():
    """Verifica se o Poppler está instalado e acessível"""
    try:
        subprocess.run(["pdfinfo", "-v"], capture_output=True, check=True)
        return True
    except Exception as e:
        st.error(f"ERRO CRÍTICO: Poppler não está instalado corretamente. Detalhes: {str(e)}")
        st.error("Solução: Execute 'sudo apt-get install poppler-utils'")
        return False

# Configuração do Tesseract
try:
    pytesseract.pytesseract.tesseract_cmd = os.getenv('TESSERACT_CMD', '/usr/bin/tesseract')
    TESSERACT_CONFIG = "--oem 3 --psm 6 -l por+eng"
except Exception as e:
    st.error(f"Erro na configuração do Tesseract: {str(e)}")
    st.stop()

# ========== MODELOS DE DOCUMENTOS COM PADRÕES ESPECÍFICOS ========== #
DOCUMENTOS_PADRAO = [
    {
        "nome": "PORTARIA DA SINDICÂNCIA ESPECIAL",
        "artigo": "NI 1.26 Art. 5º",
        "padroes_texto": [
            r"PORTARIA\s+N[º°]\s*\d+/SINDASV/\d{4}",
            r"INSTAURAÇÃO\s+DE\s+SINDICÂNCIA\s+ESPECIAL",
            r"PORTARIA\s+DE\s+SINDICÂNCIA\s+ESPECIAL"
        ],
        "palavras_chave": ["portaria", "sindicância", "especial", "instauração"],
        "pagina_referencia": 3
    },
    # Adicione outros documentos conforme necessário
]

# ========== FUNÇÕES DE PROCESSAMENTO ========== #
def pagina_vazia(img, threshold=0.95):
    """Verifica se a página é predominantemente vazia"""
    try:
        img_np = np.array(img.convert('L'))
        white_pixels = np.sum(img_np > 200)
        total_pixels = img_np.size
        return (white_pixels / total_pixels) > threshold
    except Exception as e:
        st.warning(f"Erro ao verificar página vazia: {str(e)}")
        return False

def processar_imagem_ocr(img):
    """Processa imagem para OCR com tratamento de erro"""
    try:
        img = img.convert('L')  # Converte para escala de cinza
        return pytesseract.image_to_string(img, config=TESSERACT_CONFIG)
    except Exception as e:
        st.error(f"Erro no OCR: {str(e)}")
        return ""

def extrair_texto_pdf(uploaded_file, modo_rapido=False):
    """Extrai texto do PDF com registro das páginas"""
    try:
        if not verificar_poppler():
            return None

        dpi = 200 if modo_rapido else 300
        max_pages = 10 if modo_rapido else None
        
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()
        
        imagens = convert_from_bytes(
            file_bytes,
            dpi=dpi,
            first_page=1,
            last_page=max_pages,
            poppler_path="/usr/bin"
        )
        
        textos_paginas = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, img in enumerate(imagens):
            status_text.text(f"Processando página {i+1}/{len(imagens)}...")
            if not pagina_vazia(img):
                texto = processar_imagem_ocr(img)
                textos_paginas.append((i+1, texto.upper()))
            progress_bar.progress((i + 1) / len(imagens))
        
        status_text.empty()
        progress_bar.empty()
        return textos_paginas

    except Exception as e:
        st.error(f"Falha ao processar PDF: {str(e)}")
        return None

def identificar_documentos(textos_paginas):
    """Identifica documentos com registro das páginas onde foram encontrados"""
    resultados = {}
    
    for doc in DOCUMENTOS_PADRAO:
        ocorrencias = []
        for page_num, text in textos_paginas:
            encontrado = False
            for padrao in doc["padroes_texto"]:
                if re.search(padrao, text, re.IGNORECASE):
                    ocorrencias.append(page_num)
                    encontrado = True
                    break
            
            if not encontrado and any(palavra.lower() in text.lower() for palavra in doc["palavras_chave"]):
                ocorrencias.append(page_num)
        
        if ocorrencias:
            resultados[doc["nome"]] = {
                "artigo": doc["artigo"],
                "paginas": ocorrencias,
                "pagina_referencia": doc["pagina_referencia"]
            }
    
    # DEBUG: Mostra dados identificados
    st.subheader("Debug - Dados Identificados")
    st.write(resultados)
    
    return resultados

def formatar_paginas(paginas):
    """Formata números de páginas em intervalos (ex: [1,2,3,5] → '1-3, 5')"""
    if not paginas:
        return ""
    
    paginas = sorted(set(paginas))
    ranges = []
    start = paginas[0]
    
    for i in range(1, len(paginas)):
        if paginas[i] != paginas[i-1] + 1:
            if start == paginas[i-1]:
                ranges.append(str(start))
            else:
                ranges.append(f"{start}-{paginas[i-1]}")
            start = paginas[i]
    
    if start == paginas[-1]:
        ranges.append(str(start))
    else:
        ranges.append(f"{start}-{paginas[-1]}")
    
    return ", ".join(ranges)

def gerar_relatorio(documentos_identificados):
    """Gera relatório em DOCX com formatação aprimorada"""
    try:
        doc = Document()
        
        # Cabeçalho
        title = doc.add_heading('RELATÓRIO DE ANÁLISE DOCUMENTAL', level=1)
        title.alignment = 1  # Centralizado
        
        # Data da análise
        doc.add_paragraph(f"Data da análise: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()

        # Documentos identificados
        doc.add_heading('DOCUMENTOS IDENTIFICADOS', level=2)
        if documentos_identificados:
            for doc_name, info in documentos_identificados.items():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"✓ {doc_name}").bold = True
                p.add_run(f" (Art. {info['artigo']})")
                p.add_run(f" - Páginas: {formatar_paginas(info['paginas'])}").italic = True
                p.add_run(f" | Pg. referência: {info['pagina_referencia']}")
        else:
            doc.add_paragraph("Nenhum documento padrão identificado", style='List Bullet')

        # Documentos faltantes
        doc.add_heading('DOCUMENTOS FALTANTES', level=2)
        for doc_padrao in DOCUMENTOS_PADRAO:
            if doc_padrao["nome"] not in documentos_identificados:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"✗ {doc_padrao['nome']}").bold = True
                p.add_run(f" (Art. {doc_padrao['artigo']})")
                p.add_run(f" - Pg. referência: {doc_padrao['pagina_referencia']}").italic = True

        # Rodapé
        doc.add_page_break()
        footer = doc.add_paragraph("BM/RS - Seção de Afastamentos e Acidentes")
        footer.alignment = 1  # Centralizado

        # Salva em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        st.error(f"Erro ao gerar relatório: {str(e)}")
        return None

# ========== INTERFACE PRINCIPAL ========== #
def main():
    st.title("🛡️ Sistema de Análise Documental - BM/RS")
    st.markdown("### Análise de Documentos com Referência de Páginas")
    
    # Verificação inicial do ambiente
    if not verificar_poppler():
        st.stop()
    
    # Upload do arquivo
    uploaded_file = st.file_uploader(
        "Carregue o arquivo PDF para análise",
        type=["pdf"],
        accept_multiple_files=False
    )
    
    # Opções de processamento
    modo_rapido = st.checkbox("Modo rápido (analisa apenas as primeiras 10 páginas)", value=True)
    
    # Botão de análise
    if uploaded_file and st.button("Iniciar Análise"):
        with st.spinner("Processando documento..."):
            # Extrai texto do PDF com numeração de páginas
            textos_paginas = extrair_texto_pdf(uploaded_file, modo_rapido)
            
            if textos_paginas:
                # DEBUG: Mostra texto extraído (3 primeiras páginas)
                st.subheader("Debug - Texto Extraído (Amostra)")
                for pagina, texto in textos_paginas[:3]:
                    st.write(f"Página {pagina} (primeiros 200 caracteres):")
                    st.text(texto[:200] + "...")
                
                # Identifica documentos
                documentos_identificados = identificar_documentos(textos_paginas)
                
                # Armazena resultados
                st.session_state.resultados = {
                    "documentos_identificados": documentos_identificados,
                    "textos_paginas": textos_paginas
                }
    
    # Exibição de resultados
    if 'resultados' in st.session_state:
        st.success("Análise concluída com sucesso!")
        
        # Mostrar documentos identificados
        if st.session_state.resultados["documentos_identificados"]:
            st.subheader("Documentos Identificados")
            for doc_name, info in st.session_state.resultados["documentos_identificados"].items():
                st.success(f"**{doc_name}** (Art. {info['artigo']})")
                st.write(f"🔍 Encontrado nas páginas: {formatar_paginas(info['paginas'])}")
                st.write(f"📌 Página de referência: {info['pagina_referencia']}")
        else:
            st.warning("Nenhum documento padrão foi identificado")
        
        # Botão para download do relatório
        relatorio = gerar_relatorio(st.session_state.resultados["documentos_identificados"])
        
        if relatorio:
            st.download_button(
                label="📄 Baixar Relatório Completo (DOCX)",
                data=relatorio,
                file_name=f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
